
import docx

doc = docx.Document(r'd:\proyecto\Plan_Trabajo_Concurso_Jefatura_Software.docx')
fullText = []
for para in doc.paragraphs:
    if para.text.strip():
        fullText.append(para.text)
for table in doc.tables:
    for row in table.rows:
        row_txt = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
        if row_txt:
            fullText.append(row_txt)

extracted = "\n".join(fullText)
with open(r'd:\proyecto\scratch\guido_project_text.txt', 'w', encoding='utf8') as f:
    f.write(extracted)

print(f"Extracted {len(fullText)} paragraphs/table rows. Total characters: {len(extracted)}")
print("\n=== FIRST 1500 CHARACTERS OF GUIDO PROJECT ===")
print(extracted[:1500])
